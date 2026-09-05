from __future__ import annotations

import re
import zipfile
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from xml.etree import ElementTree


@dataclass
class ParsedParagraph:
    text: str
    kind: str = "paragraph"


@dataclass
class ParsedChapter:
    title: str
    paragraphs: list[ParsedParagraph] = field(default_factory=list)


def _clean(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\u00a0", " ")).strip()


def _is_chapter_heading(text: str, language: str) -> bool:
    if not text or len(text) > 140:
        return False
    if re.search(r"^(chapter|part|section)\s+[\w\d-]+", text, re.I):
        return True
    if language == "zh" and re.search(r"^(第[一二三四五六七八九十百千万0-9]+[章节篇部]|卷\s*[一二三四五六七八九十0-9]+)", text):
        return True
    return text.isupper() and len(text.split()) <= 12


def _from_lines(lines: list[str], language: str, default_title: str) -> list[ParsedChapter]:
    chapters: list[ParsedChapter] = []
    current: ParsedChapter | None = None
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        text = _clean(" ".join(paragraph_lines))
        if text and current is not None:
            current.paragraphs.append(ParsedParagraph(text))
        paragraph_lines = []

    def start_chapter(title: str) -> None:
        nonlocal current
        flush_paragraph()
        current = ParsedChapter(_clean(title) or f"{default_title} {len(chapters) + 1}")
        chapters.append(current)

    start_chapter(default_title)
    for raw in lines:
        line = raw.strip()
        if not line:
            flush_paragraph()
            continue
        if _is_chapter_heading(line, language):
            start_chapter(line)
            continue
        paragraph_lines.append(line)
    flush_paragraph()
    chapters = [chapter for chapter in chapters if chapter.paragraphs or len(chapters) == 1]
    if not chapters:
        chapters = [ParsedChapter(default_title)]
    return chapters


def parse_pdf(path: Path, title: str) -> list[ParsedChapter]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("缺少 PDF 解析依赖，请安装 pypdf") from exc

    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "") for page in reader.pages]
    repeated: dict[str, int] = {}
    for page in pages:
        lines = [line.strip() for line in page.splitlines() if line.strip()]
        for line in set(lines[:2] + lines[-2:]):
            if len(line) < 100:
                repeated[line] = repeated.get(line, 0) + 1
    threshold = max(3, int(len(pages) * 0.3))
    boilerplate = {line for line, count in repeated.items() if count >= threshold}

    lines: list[str] = []
    for page in pages:
        page_lines = [line.rstrip() for line in page.splitlines()]
        while page_lines and (not page_lines[0].strip() or page_lines[0].strip() in boilerplate):
            page_lines.pop(0)
        while page_lines and (not page_lines[-1].strip() or page_lines[-1].strip() in boilerplate):
            page_lines.pop()
        for line in page_lines:
            clean = line.strip()
            if re.fullmatch(r"[-–—]?\s*\d+\s*[-–—]?", clean):
                continue
            if clean in boilerplate:
                continue
            # Extracted PDF text often uses a soft line break within a word.
            if lines and clean and not lines[-1].endswith((".", "?", "!", ":", ";", '”', '"', "。", "！", "？", "：", "；")) and not clean[:1].isupper():
                lines[-1] = f"{lines[-1]} {clean}"
            else:
                lines.append(line)
        lines.append("")
    metadata_title = (reader.metadata.title if reader.metadata else None) or title
    return _from_lines(lines, "en", _clean(metadata_title) or title)


class _EpubParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[tuple[str, str]] = []
        self._tag: str | None = None
        self._parts: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in {"p", "h1", "h2", "h3", "h4", "h5", "h6", "blockquote", "li"}:
            self._tag = tag
            self._parts = []

    def handle_data(self, data: str) -> None:
        if self._tag:
            self._parts.append(data)

    def handle_endtag(self, tag: str) -> None:
        if self._tag == tag.lower():
            text = _clean("".join(self._parts))
            if text:
                self.blocks.append((self._tag, text))
            self._tag = None
            self._parts = []


def _epub_spine(zf: zipfile.ZipFile) -> list[str]:
    try:
        container = ElementTree.fromstring(zf.read("META-INF/container.xml"))
        rootfile = next(iter(container.iter()))
        rootfile_path = next((node.attrib.get("full-path") for node in container.iter() if node.tag.endswith("rootfile")), "")
        opf = ElementTree.fromstring(zf.read(rootfile_path))
        base = str(Path(rootfile_path).parent)
        manifest = {item.attrib["id"]: item.attrib["href"] for item in opf.iter() if item.tag.endswith("item")}
        files = [manifest[item.attrib["id"]] for item in opf.iter() if item.tag.endswith("itemref") and item.attrib.get("idref") in manifest]
        return [str(Path(base) / item) if base != "." else item for item in files]
    except Exception:
        return sorted(name for name in zf.namelist() if name.lower().endswith((".xhtml", ".html", ".htm")))


def parse_epub(path: Path, title: str) -> list[ParsedChapter]:
    with zipfile.ZipFile(path) as zf:
        chapters: list[ParsedChapter] = []
        for filename in _epub_spine(zf):
            try:
                parser = _EpubParser()
                parser.feed(zf.read(filename).decode("utf-8", errors="ignore"))
            except KeyError:
                continue
            if not parser.blocks:
                continue
            chapter_headings = [text for tag, text in parser.blocks if tag.startswith("h")]
            chapter_title = chapter_headings[0] if chapter_headings else f"{title} {len(chapters) + 1}"
            chapter = ParsedChapter(chapter_title)
            for tag, text in parser.blocks:
                if tag.startswith("h"):
                    continue
                chapter.paragraphs.append(ParsedParagraph(text))
            if chapter.paragraphs:
                chapters.append(chapter)
        return chapters or [ParsedChapter(title)]


def parse_docx(path: Path, title: str) -> list[ParsedChapter]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 DOCX 解析依赖，请安装 python-docx") from exc

    document = Document(str(path))
    chapters: list[ParsedChapter] = []
    current = ParsedChapter(title)
    chapters.append(current)
    for paragraph in document.paragraphs:
        text = _clean(paragraph.text)
        if not text:
            continue
        style = (paragraph.style.name or "").lower() if paragraph.style else ""
        if style.startswith("heading") or _is_chapter_heading(text, "zh"):
            current = ParsedChapter(text)
            chapters.append(current)
        else:
            current.paragraphs.append(ParsedParagraph(text))
    return [chapter for chapter in chapters if chapter.paragraphs] or [ParsedChapter(title)]


def parse_source(path: Path, language: str, original_name: str) -> tuple[str, list[ParsedChapter]]:
    suffix = path.suffix.lower()
    title = Path(original_name).stem.replace("_", " ").strip() or "未命名书籍"
    if suffix == ".pdf":
        return title, parse_pdf(path, title)
    if suffix == ".epub":
        return title, parse_epub(path, title)
    if suffix == ".docx":
        return title, parse_docx(path, title)
    expected = "PDF" if language == "en" else "EPUB 或 DOCX"
    raise ValueError(f"{language} 文件必须是 {expected}")
